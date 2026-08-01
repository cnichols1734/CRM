"""Attachment parsing and intent classification for UI B.O.B. chat.

Uploads are treated as untrusted data. This module classifies what kind of
file arrived, extracts a bounded representation, and decides whether the
user asked to analyze it or import contacts from it.
"""
from __future__ import annotations

import csv
import io
import logging
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any, Optional

from services.inbound_payload import image_to_base64_jpeg

logger = logging.getLogger(__name__)

# Hard ceilings for B.O.B. attachment processing.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_SPREADSHEET_ROWS = 500
MAX_SPREADSHEET_COLUMNS = 50
MAX_SPREADSHEET_CELLS = 25_000
MAX_TEXT_CHARS = 60_000
MAX_PDF_TEXT_PAGES = 50
MAX_PDF_RENDER_PAGES = 10
MAX_DOC_CANDIDATES = 25
MAX_IMAGE_CANDIDATES = 5

INTENT_ANALYZE = 'analyze'
INTENT_CREATE = 'create_contacts'
INTENT_AMBIGUOUS = 'ambiguous'

KIND_IMAGE = 'image'
KIND_CSV = 'csv'
KIND_XLSX = 'xlsx'
KIND_XLS = 'xls'
KIND_TXT = 'txt'
KIND_VCF = 'vcf'
KIND_PDF = 'pdf'
KIND_DOCX = 'docx'
KIND_UNSUPPORTED = 'unsupported'

IMAGE_MIMES = {
    'image/jpeg', 'image/jpg', 'image/png', 'image/gif', 'image/webp',
}
CSV_MIMES = {'text/csv', 'application/csv'}
XLSX_MIMES = {
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
}
XLS_MIMES = {'application/vnd.ms-excel'}
TXT_MIMES = {'text/plain'}
VCF_MIMES = {'text/vcard', 'text/x-vcard', 'text/directory'}
PDF_MIMES = {'application/pdf'}
DOCX_MIMES = {
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
}
# Legacy Word is intentionally unsupported.
DOC_MIMES = {'application/msword'}

_CREATE_INTENT_RE = re.compile(
    r'(?i)\b('
    r'create(\s+(this|the|a|my|these|those))?\s+contacts?'
    r'|add(\s+(this|the|a|my|these|those))?\s+contacts?'
    r'|save(\s+(this|the|a|my|these|those))?\s+contacts?'
    r'|import(\s+(this|the|these|those|my))?\s+(contacts?|list|file|spreadsheet|csv|excel)?'
    r'|add(\s+(them|him|her|this|these|those))?\s+to\s+(my\s+)?(crm|contacts?)'
    r'|save(\s+(them|him|her|this|these|those))?\s+to\s+(my\s+)?(crm|contacts?)'
    r'|put\s+(this|them|him|her|these|those)\s+in\s+(my\s+)?(crm|contacts?)'
    r'|load(\s+(this|the|these|those))?\s+(list|contacts?|file|spreadsheet|csv|excel)'
    r'|new\s+contact'
    r')\b'
)

_ANALYZE_INTENT_RE = re.compile(
    r'(?i)\b('
    r'what|who|where|when|why|how|summar(y|ize)|analyse|analyze|explain|'
    r'describe|tell me|look at|review|read|count|how many|find|show|'
    r'compare|average|total|missing|duplicate|filter'
    r')\b'
)

_ACTIONISH_RE = re.compile(
    r'(?i)\b(add|create|save|import|load|put|upload|use this|from this)\b'
)


class AttachmentParseError(Exception):
    def __init__(self, message: str):
        super().__init__(message)
        self.message = message


@dataclass
class ParsedAttachment:
    kind: str
    filename: str
    mime: str
    size: int
    text: str = ''
    rows: list[dict[str, Any]] = field(default_factory=list)
    headers: list[str] = field(default_factory=list)
    image_jpeg_b64: str | None = None
    pdf_page_images: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    truncated: bool = False
    stats: dict[str, Any] = field(default_factory=dict)

    @property
    def is_tabular(self) -> bool:
        return self.kind in {KIND_CSV, KIND_XLSX, KIND_XLS}

    @property
    def is_empty_content(self) -> bool:
        if self.kind == KIND_IMAGE:
            return not self.image_jpeg_b64
        if self.is_tabular:
            return not self.rows
        return not (self.text or '').strip() and not self.pdf_page_images


def classify_kind(filename: str, mime: str) -> str:
    name = (filename or '').lower()
    mime = (mime or '').lower()

    if mime in DOC_MIMES or name.endswith('.doc'):
        return KIND_UNSUPPORTED
    if mime in IMAGE_MIMES or any(name.endswith(ext) for ext in (
            '.jpg', '.jpeg', '.png', '.gif', '.webp')):
        return KIND_IMAGE
    if name.endswith('.csv') or mime in CSV_MIMES:
        return KIND_CSV
    if name.endswith('.xlsx') or mime in XLSX_MIMES:
        return KIND_XLSX
    if name.endswith('.xls') or (
            mime in XLS_MIMES and not name.endswith('.xlsx')
            and not name.endswith('.csv')):
        # Outlook sometimes labels CSV as ms-excel; prefer extension.
        return KIND_XLS
    if name.endswith('.vcf') or mime in VCF_MIMES:
        return KIND_VCF
    if name.endswith('.txt') or mime in TXT_MIMES:
        return KIND_TXT
    if name.endswith('.pdf') or mime in PDF_MIMES:
        return KIND_PDF
    if name.endswith('.docx') or mime in DOCX_MIMES:
        return KIND_DOCX
    return KIND_UNSUPPORTED


def classify_attachment_intent(
    message: str,
    attachment_kind: str,
    *,
    is_empty: bool = False,
) -> str:
    """Return analyze | create_contacts | ambiguous.

    Strong deterministic phrases win. Empty spreadsheets/documents ask.
    Empty images stay on the analyze path so vision can still describe them.
    """
    text = (message or '').strip()

    if is_empty and attachment_kind != KIND_IMAGE:
        return INTENT_AMBIGUOUS

    if not text:
        if attachment_kind == KIND_IMAGE:
            return INTENT_ANALYZE
        return INTENT_AMBIGUOUS

    if _CREATE_INTENT_RE.search(text):
        return INTENT_CREATE
    if _ANALYZE_INTENT_RE.search(text) and not _CREATE_INTENT_RE.search(text):
        return INTENT_ANALYZE
    if _ACTIONISH_RE.search(text):
        return _classify_intent_with_model(text, attachment_kind)
    return INTENT_ANALYZE


def _classify_intent_with_model(text: str, attachment_kind: str) -> str:
    """Small structured classifier for ambiguous action-like phrasing."""
    try:
        import json
        import openai
        from config import Config

        if not Config.OPENAI_API_KEY:
            return INTENT_AMBIGUOUS

        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model='gpt-4.1-mini',
            messages=[
                {
                    'role': 'system',
                    'content': (
                        'Classify the user intent for a CRM chat attachment. '
                        'Return JSON only with key "intent" equal to one of: '
                        'analyze, create_contacts, ambiguous. '
                        'create_contacts means they want people saved into the CRM. '
                        'analyze means they want questions answered about the file. '
                        'ambiguous means you cannot tell.'
                    ),
                },
                {
                    'role': 'user',
                    'content': (
                        f'Attachment kind: {attachment_kind}\n'
                        f'User message: {text[:1000]}'
                    ),
                },
            ],
            response_format={'type': 'json_object'},
            temperature=0,
        )
        raw = response.choices[0].message.content or '{}'
        parsed = json.loads(raw)
        intent = (parsed.get('intent') or '').strip()
        if intent in {INTENT_ANALYZE, INTENT_CREATE, INTENT_AMBIGUOUS}:
            return intent
    except Exception:
        logger.warning('Attachment intent classifier failed', exc_info=True)
    return INTENT_AMBIGUOUS


def parse_attachment(
    data: bytes,
    *,
    filename: str,
    mime: str,
    max_rows: int = MAX_SPREADSHEET_ROWS,
    max_upload_bytes: int = MAX_UPLOAD_BYTES,
) -> ParsedAttachment:
    if data is None:
        raise AttachmentParseError('That attachment looked empty.')
    if len(data) > max_upload_bytes:
        raise AttachmentParseError(
            f'That file is larger than {max_upload_bytes // (1024 * 1024)}MB.'
        )

    kind = classify_kind(filename, mime)
    if kind == KIND_UNSUPPORTED:
        raise AttachmentParseError(
            'That file type is not supported in chat. '
            'Use an image, CSV, Excel (.xlsx/.xls), TXT, VCF, PDF, or DOCX.'
        )

    parsed = ParsedAttachment(
        kind=kind,
        filename=filename or 'attachment',
        mime=mime or 'application/octet-stream',
        size=len(data),
    )

    if kind == KIND_IMAGE:
        _parse_image(parsed, data)
    elif kind == KIND_CSV:
        _parse_csv(parsed, data, max_rows=max_rows)
    elif kind == KIND_XLSX:
        _parse_xlsx(parsed, data, max_rows=max_rows)
    elif kind == KIND_XLS:
        _parse_xls(parsed, data, max_rows=max_rows)
    elif kind in {KIND_TXT, KIND_VCF}:
        _parse_text(parsed, data)
    elif kind == KIND_PDF:
        _parse_pdf(parsed, data)
    elif kind == KIND_DOCX:
        _parse_docx(parsed, data)

    if parsed.is_tabular:
        parsed.stats = compute_tabular_stats(parsed.rows, parsed.headers)
    return parsed


def compute_tabular_stats(
    rows: list[dict[str, Any]],
    headers: list[str],
) -> dict[str, Any]:
    missing: dict[str, int] = {h: 0 for h in headers}
    for row in rows:
        for h in headers:
            value = row.get(h)
            if value in (None, ''):
                missing[h] += 1

    seen_emails: set[str] = set()
    email_dupes = 0
    for row in rows:
        email = str(row.get('email') or row.get('Email') or '').strip().lower()
        if not email:
            continue
        if email in seen_emails:
            email_dupes += 1
        else:
            seen_emails.add(email)

    return {
        'row_count': len(rows),
        'column_count': len(headers),
        'columns': headers,
        'missing_by_column': missing,
        'duplicate_email_rows': email_dupes,
        'sample_rows': rows[:5],
    }


def query_tabular(
    rows: list[dict[str, Any]],
    *,
    operation: str = 'summary',
    column: str | None = None,
    filters: list[dict] | None = None,
    limit: int = 20,
) -> dict[str, Any]:
    """Deterministic spreadsheet operations for inspect_attachment."""
    working = list(rows)
    for rule in filters or []:
        if not isinstance(rule, dict):
            continue
        col = rule.get('column')
        op = (rule.get('op') or 'eq').lower()
        expected = rule.get('value')
        if not col:
            continue
        next_rows = []
        for row in working:
            value = row.get(col)
            if op == 'eq' and str(value) == str(expected):
                next_rows.append(row)
            elif op == 'contains' and expected is not None and str(expected).lower() in str(value or '').lower():
                next_rows.append(row)
            elif op == 'empty' and value in (None, ''):
                next_rows.append(row)
            elif op == 'not_empty' and value not in (None, ''):
                next_rows.append(row)
        working = next_rows

    op = (operation or 'summary').lower()
    if op == 'count':
        return {'operation': 'count', 'count': len(working)}
    if op == 'sample':
        return {
            'operation': 'sample',
            'count': len(working),
            'rows': working[: max(1, min(limit, 50))],
        }
    if op in {'sum', 'min', 'max', 'average'} and column:
        nums = []
        for row in working:
            try:
                nums.append(float(str(row.get(column)).replace(',', '')))
            except (TypeError, ValueError):
                continue
        if not nums:
            return {'operation': op, 'column': column, 'count': 0, 'value': None}
        if op == 'sum':
            value = sum(nums)
        elif op == 'min':
            value = min(nums)
        elif op == 'max':
            value = max(nums)
        else:
            value = sum(nums) / len(nums)
        return {
            'operation': op,
            'column': column,
            'count': len(nums),
            'value': value,
        }

    headers = list(rows[0].keys()) if rows else []
    return {
        'operation': 'summary',
        'stats': compute_tabular_stats(working, headers),
    }


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def _decode_text(data: bytes) -> str:
    for enc in ('utf-8-sig', 'utf-8', 'latin-1'):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


def _parse_image(parsed: ParsedAttachment, data: bytes) -> None:
    b64 = image_to_base64_jpeg(data)
    if not b64:
        raise AttachmentParseError(
            "Couldn't read that image. Try a clearer photo or a different format."
        )
    parsed.image_jpeg_b64 = b64


def _parse_text(parsed: ParsedAttachment, data: bytes) -> None:
    text = _decode_text(data)
    if len(text) > MAX_TEXT_CHARS:
        parsed.text = text[:MAX_TEXT_CHARS]
        parsed.truncated = True
        parsed.warnings.append(
            f'Text truncated to the first {MAX_TEXT_CHARS:,} characters.'
        )
    else:
        parsed.text = text


def _parse_csv(
    parsed: ParsedAttachment,
    data: bytes,
    *,
    max_rows: int = MAX_SPREADSHEET_ROWS,
) -> None:
    text = _decode_text(data)
    reader = csv.DictReader(io.StringIO(text))
    headers = list(reader.fieldnames or [])
    if headers and headers[0].startswith('\ufeff'):
        headers[0] = headers[0].replace('\ufeff', '')
    if len(headers) > MAX_SPREADSHEET_COLUMNS:
        raise AttachmentParseError(
            f'That spreadsheet has more than {MAX_SPREADSHEET_COLUMNS} columns.'
        )
    rows = []
    for idx, row in enumerate(reader, start=1):
        if idx > max_rows:
            parsed.truncated = True
            parsed.warnings.append(
                f'Only the first {max_rows} rows were loaded. '
                'Use Contacts → Import for larger files.'
            )
            break
        cleaned = {
            (k or '').strip(): ('' if v is None else str(v).strip())
            for k, v in row.items()
        }
        rows.append(cleaned)
        if len(rows) * max(len(headers), 1) > MAX_SPREADSHEET_CELLS:
            parsed.truncated = True
            parsed.warnings.append(
                'Spreadsheet cell limit reached; remaining rows were skipped.'
            )
            break
    parsed.headers = [h.strip() for h in headers]
    parsed.rows = rows
    parsed.text = (
        f'CSV with {len(rows)} data row(s) and columns: '
        + ', '.join(parsed.headers[:20])
    )


def _guard_zip_bomb(data: bytes) -> None:
    """Reject pathological Office Open XML packages before full parse."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            total_uncompressed = 0
            for info in zf.infolist():
                total_uncompressed += int(info.file_size or 0)
                if info.file_size and info.compress_size:
                    ratio = info.file_size / max(info.compress_size, 1)
                    if ratio > 200 and info.file_size > 5 * 1024 * 1024:
                        raise AttachmentParseError(
                            'That Excel file looks unsafe and was rejected.'
                        )
                if total_uncompressed > 80 * 1024 * 1024:
                    raise AttachmentParseError(
                        'That Excel file expands too large and was rejected.'
                    )
    except zipfile.BadZipFile as exc:
        raise AttachmentParseError(
            'That Excel file could not be opened.'
        ) from exc


def _parse_xlsx(
    parsed: ParsedAttachment,
    data: bytes,
    *,
    max_rows: int = MAX_SPREADSHEET_ROWS,
) -> None:
    if data[:2] != b'PK':
        raise AttachmentParseError('That does not look like an .xlsx file.')
    _guard_zip_bomb(data)
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise AttachmentParseError('Excel support is unavailable right now.') from exc

    try:
        wb = load_workbook(
            io.BytesIO(data),
            read_only=True,
            data_only=True,
            keep_links=False,
        )
    except Exception as exc:
        raise AttachmentParseError(
            'That Excel file could not be opened.'
        ) from exc

    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            parsed.headers = []
            parsed.rows = []
            return

        headers = []
        for idx, cell in enumerate(header_row):
            if idx >= MAX_SPREADSHEET_COLUMNS:
                raise AttachmentParseError(
                    f'That spreadsheet has more than {MAX_SPREADSHEET_COLUMNS} columns.'
                )
            name = str(cell).strip() if cell is not None else f'column_{idx + 1}'
            headers.append(name or f'column_{idx + 1}')

        rows = []
        for idx, values in enumerate(rows_iter, start=1):
            if idx > max_rows:
                parsed.truncated = True
                parsed.warnings.append(
                    f'Only the first {max_rows} rows were loaded. '
                    'Use Contacts → Import for larger files.'
                )
                break
            if values is None or all(v is None or str(v).strip() == '' for v in values):
                continue
            row = {}
            for col_idx, header in enumerate(headers):
                value = values[col_idx] if col_idx < len(values) else None
                row[header] = '' if value is None else str(value).strip()
            rows.append(row)
            if len(rows) * max(len(headers), 1) > MAX_SPREADSHEET_CELLS:
                parsed.truncated = True
                parsed.warnings.append(
                    'Spreadsheet cell limit reached; remaining rows were skipped.'
                )
                break
        parsed.headers = headers
        parsed.rows = rows
        parsed.text = (
            f'Excel sheet with {len(rows)} data row(s) and columns: '
            + ', '.join(headers[:20])
        )
    finally:
        wb.close()


def _parse_xls(
    parsed: ParsedAttachment,
    data: bytes,
    *,
    max_rows: int = MAX_SPREADSHEET_ROWS,
) -> None:
    try:
        import xlrd
    except ImportError as exc:
        raise AttachmentParseError('Legacy Excel support is unavailable.') from exc

    try:
        book = xlrd.open_workbook(file_contents=data)
    except Exception as exc:
        raise AttachmentParseError(
            'That .xls file could not be opened.'
        ) from exc

    sheet = book.sheet_by_index(0)
    if sheet.ncols > MAX_SPREADSHEET_COLUMNS:
        raise AttachmentParseError(
            f'That spreadsheet has more than {MAX_SPREADSHEET_COLUMNS} columns.'
        )
    if sheet.nrows == 0:
        parsed.headers = []
        parsed.rows = []
        return

    headers = []
    for col in range(sheet.ncols):
        value = sheet.cell_value(0, col)
        name = str(value).strip() if value not in (None, '') else f'column_{col + 1}'
        headers.append(name)

    rows = []
    max_data_rows = min(sheet.nrows - 1, max_rows)
    if sheet.nrows - 1 > max_rows:
        parsed.truncated = True
        parsed.warnings.append(
            f'Only the first {max_rows} rows were loaded. '
            'Use Contacts → Import for larger files.'
        )
    for row_idx in range(1, max_data_rows + 1):
        row = {}
        empty = True
        for col, header in enumerate(headers):
            value = sheet.cell_value(row_idx, col)
            text = '' if value in (None, '') else str(value).strip()
            if text:
                empty = False
            row[header] = text
        if empty:
            continue
        rows.append(row)
        if len(rows) * max(len(headers), 1) > MAX_SPREADSHEET_CELLS:
            parsed.truncated = True
            parsed.warnings.append(
                'Spreadsheet cell limit reached; remaining rows were skipped.'
            )
            break

    parsed.headers = headers
    parsed.rows = rows
    parsed.text = (
        f'Excel sheet with {len(rows)} data row(s) and columns: '
        + ', '.join(headers[:20])
    )


def _parse_pdf(parsed: ParsedAttachment, data: bytes) -> None:
    try:
        import fitz
    except ImportError as exc:
        raise AttachmentParseError('PDF support is unavailable right now.') from exc

    try:
        doc = fitz.open(stream=data, filetype='pdf')
    except Exception as exc:
        raise AttachmentParseError('That PDF could not be opened.') from exc

    try:
        if doc.is_encrypted:
            raise AttachmentParseError(
                'That PDF is encrypted and cannot be read in chat.'
            )

        chunks = []
        page_count = doc.page_count
        text_pages = min(page_count, MAX_PDF_TEXT_PAGES)
        for index in range(text_pages):
            page_text = (doc.load_page(index).get_text('text') or '').strip()
            if page_text:
                chunks.append(f'--- Page {index + 1} ---\n{page_text}')
        text = '\n\n'.join(chunks)
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS]
            parsed.truncated = True
            parsed.warnings.append(
                f'PDF text truncated to the first {MAX_TEXT_CHARS:,} characters.'
            )
        if page_count > MAX_PDF_TEXT_PAGES:
            parsed.truncated = True
            parsed.warnings.append(
                f'Only the first {MAX_PDF_TEXT_PAGES} PDF pages were read for text.'
            )
        parsed.text = text

        # Scanned PDFs: render a few pages for vision.
        if not text.strip():
            images = []
            render_pages = min(page_count, MAX_PDF_RENDER_PAGES)
            for index in range(render_pages):
                pix = doc.load_page(index).get_pixmap(dpi=150)
                png_bytes = pix.tobytes('png')
                b64 = image_to_base64_jpeg(png_bytes)
                if b64:
                    images.append(b64)
            parsed.pdf_page_images = images
            if images:
                parsed.warnings.append(
                    'No selectable text found; used page images instead.'
                )
            if page_count > MAX_PDF_RENDER_PAGES:
                parsed.truncated = True
                parsed.warnings.append(
                    f'Only the first {MAX_PDF_RENDER_PAGES} pages were rendered.'
                )
    finally:
        doc.close()


def _parse_docx(parsed: ParsedAttachment, data: bytes) -> None:
    if data[:2] != b'PK':
        raise AttachmentParseError('That does not look like a .docx file.')
    _guard_zip_bomb(data)
    try:
        from docx import Document
    except ImportError as exc:
        raise AttachmentParseError('DOCX support is unavailable right now.') from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise AttachmentParseError('That Word document could not be opened.') from exc

    parts = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or '').strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [
                (cell.text or '').strip()
                for cell in row.cells
                if (cell.text or '').strip()
            ]
            if cells:
                parts.append(' | '.join(cells))

    text = '\n'.join(parts)
    if len(text) > MAX_TEXT_CHARS:
        parsed.text = text[:MAX_TEXT_CHARS]
        parsed.truncated = True
        parsed.warnings.append(
            f'Document text truncated to the first {MAX_TEXT_CHARS:,} characters.'
        )
    else:
        parsed.text = text


def extract_contact_candidates_from_attachment(
    parsed: ParsedAttachment,
    *,
    user,
    caption: str = '',
) -> list[dict]:
    """Run Magic Inbox extraction for non-tabular create-intent attachments."""
    from services.ai_service import generate_contact_extraction
    from services.messaging import photo_contacts as photo_mod

    image_blocks = []
    if parsed.image_jpeg_b64:
        image_blocks.append(parsed.image_jpeg_b64)
    image_blocks.extend(parsed.pdf_page_images[:MAX_PDF_RENDER_PAGES])

    text = caption or ''
    if parsed.text:
        text = (text + '\n\n' + parsed.text).strip()

    result = generate_contact_extraction(text=text, image_blocks=image_blocks)
    raw = result.get('contacts') or []
    if not isinstance(raw, list):
        return []

    user_email = (getattr(user, 'email', None) or '').strip().lower()
    cap = MAX_IMAGE_CANDIDATES if parsed.kind == KIND_IMAGE else MAX_DOC_CANDIDATES
    candidates = []
    for entry in raw:
        if not isinstance(entry, dict):
            continue
        normalized = photo_mod._normalize_candidate(entry, user_email=user_email)
        if normalized is not None:
            candidates.append(normalized)
        if len(candidates) >= cap:
            break
    return candidates
